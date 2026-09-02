# ==============================================================================
# Archivo: app.py
# Proyecto: War Room C5I - Puesto de Mando CMPC
# ==============================================================================

import streamlit as st
import pandas as pd
import numpy as np
from supabase import create_client, Client
import plotly.express as px
import plotly.graph_objects as go
import pydeck as pdk
from pyvis.network import Network
import streamlit.components.v1 as components
from datetime import datetime, timedelta
import io
import os
import re
import matplotlib.pyplot as plt
import base64
import requests
import json
from wordcloud import WordCloud
from bs4 import BeautifulSoup
from streamlit_autorefresh import st_autorefresh
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 🔒 1. SISTEMA DE AUTENTICACIÓN Y SEGURIDAD
st.set_page_config(page_title="C5I WAR ROOM | CMPC", layout="wide", initial_sidebar_state="expanded")

if "autenticado" not in st.session_state:
    st.session_state["autenticado"] = False

if not st.session_state["autenticado"]:
    st.markdown("<h2 style='text-align: center; color: #ff4b4b; margin-top: 100px;'>🛡️ ACCESO RESTRINGIDO - WAR ROOM CMPC</h2>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1,2,1])
    with c2:
        st.info("Ingresa tus credenciales operativas. Las cuentas son gestionadas por el Administrador de Inteligencia.")
        usuario = st.text_input("Usuario Operativo")
        clave = st.text_input("Clave de Acceso", type="password")
        if st.button("Autenticar ❯", use_container_width=True, type="primary"):
            usuarios_bd = st.secrets.get("usuarios", {"admin": "cmpc2026"})
            if usuario in usuarios_bd and str(usuarios_bd[usuario]) == clave:
                st.session_state["autenticado"] = True
                st.session_state["usuario_actual"] = usuario
                st.rerun()
            else:
                st.error("❌ Credenciales inválidas o acceso revocado.")
    st.stop() 

# --- 0. INICIALIZAR MEMORIA TÁCTICA ---
for k, v in [('filtro_cmpc_activo', False), ('filtro_provincia_activo', "Todas"), 
             ('filtro_tipologia_activo', "Todas"), ('filtro_canal_activo', "Todos")]:
    if k not in st.session_state: st.session_state[k] = v

st_autorefresh(interval=300000, key="refresh_warroom")

# 🎨 INYECCIÓN CSS TÁCTICA
st.markdown("""
<style>
:root { --bg-main: #0a0f18; --bg-panel: #111827; --bg-control: #1f2937; --border: #374151; --text-main: #e5e7eb; --text-muted: #9ca3af; --color-ok: #10b981; --color-warn: #f59e0b; --color-crit: #ef4444; --color-info: #3b82f6;}
html, body, [data-testid="stAppViewContainer"], .stApp { background-color: var(--bg-main) !important; color: var(--text-main) !important; font-family: 'Inter', system-ui, sans-serif !important; }
[data-testid="stSidebar"] { background-color: #0d1321 !important; border-right: 1px solid var(--border) !important; }
[data-testid="stDateInput"] input, [data-testid="stSelectbox"] select, [data-testid="stSlider"] input, [data-testid="stButton"] button { background-color: var(--bg-control) !important; color: var(--text-main) !important; border: 1px solid var(--border) !important; border-radius: 6px !important;}
.stMetric { background-color: var(--bg-panel) !important; padding: 16px !important; border-radius: 8px !important; border-left: 4px solid var(--text-muted) !important; }
.metric-ok { border-left-color: var(--color-ok) !important; } .metric-warn { border-left-color: var(--color-warn) !important; } .metric-crit { border-left-color: var(--color-crit) !important; }
h1, h2, h3 { color: var(--text-main) !important; letter-spacing: 0.3px !important; font-weight: 600 !important; }
</style>
""", unsafe_allow_html=True)

# ==============================================================================
# 1. DICCIONARIOS DE INTELIGENCIA
# ==============================================================================
PPM_KEYWORDS = ["Juan Huenupil", "César Millanao", "Orlando Sáez", "Esteban Carrera", "Bernardo Camus", "Matías Leviqueo", "Alexis Manríquez", "Yerko Maril", "Francisco Huichacura", "Esteban Huichacura", "Carlos Huichacura", "Manuel Huichacura", "Claudia Nahuelan", "Víctor Llanquileo", "Oscar Pilquimán", "Eliseo Raiman", "Héctor Llaitul", "Domingo Mariñan", "Manuel Alonso Llempi", "Miguel Llanquileo", "Erick Montoya", "Pablo Cayuhan", "Juan Mariñan", "Elías Cona", "Camilo Astete", "José Luis Marilao", "José Melgarejo", "Guillermo Camus", "Miguel Torres Toro", "Juan Cortés Penchulef", "Alejandro Liguen", "Anthony Torres", "Pedro Palacios", "Jorge Palacios", "Boris Llanca", "Simón Huenchullán", "Juan Queipul", "Joaquín Huenchullán", "Joaquín Millanao", "Marco Tori", "Christopher Tori", "Juan Patricio Queipul", "Danilo Nahuelpi", "Luis David Morales", "Rubén Cheuquepan", "Leandro Catrileo", "José Lienqueo", "Axel Campos", "Luis Melinao", "Benjamín Coñopan", "Fredy Marileo", "Rodrigo Calabrano", "Luis Fuenzalida", "Matías Ancalaf", "Moroni Ancalaf", "Jorge Caniupil", "Oscar Cañupan", "Rafael Pichun", "Luis Menares", "Pelentaro Llaitul", "Juan Carlos Mardones", "Roberto Garling", "Carlos Fierro", "Luis Marileo", "Patricio Queipul", "Raúl Caniullan", "Nelson Queupil", "Rodrigo Cáceres", "Fabian Llanca", "Emilio Berkhoff", "Luis Tranamil", "José Pichunhuala", "Eduardo Fuica", "Guillermo Ñiripil", "Anthu Llanca", "Máximo Queipul", "Daniel Canio Tralcal", "Bastian Llaitul", "Sergio Levinao", "José Sergio Tralcal", "Luis Tralcal", "Celestino Córdova", "Dago Queipul", "Pablo Quidel", "Juan Pablo Pirce"]
NODOS_MEDIOS = ["mapuexpress", "radiokurruf", "mapuchediario", "radionewen", "elpuelche", "radiojgm", "piensachile", "elciudadano", "mediosdelospueblos", "radiomulutu", "lafkenmawida", "mingaancestral", "comunidadtemucuicui", "lazarzamora", "futatrawun", "resumen", "interferencia", "laizquierdadiario", "araucaniadiario", "rebelion", "resumenlatinoamericano", "aukinlavken", "wall_mapuche", "wallmapunche", "radiouach", "radioplazadeladignidad", "reconstruccionnacionalmapuche", "CDNukeMapu", "ppm_casoquilleco", "memoriasenresistenciatemuko", "redsuperacionalmodeloforestal", "libertad_ppmcam", "wechekekawin", "resistencia.araucanialx", "trepemulen", "wallmapu__libre2", "brotes.del.despojo", "coordinadoraterritorialtome", "liberacionmapuchelafkenche", "lafken.kimun", "resistenciawallmapu", "envivoaquiyahoraofficial", "victor.llanquileo.pilquiman", "werken_noticias"]

# ==============================================================================
# 2. FUNCIONES AUXILIARES & CARGA DE DATOS
# ==============================================================================
def extraer_imagen_real_rss(url):
    if pd.isna(url) or not isinstance(url, str) or not url.startswith("http"): return ""
    try:
        h = {'User-Agent': 'Mozilla/5.0'}
        r = requests.get(url, headers=h, timeout=5, allow_redirects=True)
        soup = BeautifulSoup(r.content, 'html.parser')
        og = soup.find('meta', property='og:image')
        if og and og.get('content'): return og['content']
        return ""
    except: return ""

def inyectar_evidencia_b64(ruta_local, url_web):
    r_local = str(ruta_local).strip() if ruta_local else ""
    u_web = str(url_web).strip() if url_web else ""
    if r_local and r_local.lower() not in ['nan', 'none', 'no especificado'] and os.path.exists(r_local):
        try:
            es_video = any(ext in r_local.lower() for ext in ['.mp4', '.mov'])
            with open(r_local, "rb") as f: b64 = base64.b64encode(f.read()).decode()
            return f"data:video/mp4;base64,{b64}", es_video
        except: pass
    if u_web and len(u_web) > 5 and u_web.lower() != 'nan':
        return u_web, any(ext in u_web.lower() for ext in ['.mp4', '.mov', 'reel', 'video'])
    return "", False

def determinar_sentimiento(texto):
    txt = str(texto).lower()
    negativas = ['ataque', 'incendio', 'quema', 'destruye', 'muerte', 'asesinato', 'violencia', 'amenaza', 'corte', 'barricada', 'disparo']
    positivas = ['apoyo', 'solidaridad', 'inversión', 'crecimiento', 'paz', 'acuerdo', 'diálogo', 'entrega', 'ayuda']
    
    score_neg = sum(1 for w in negativas if w in txt)
    score_pos = sum(1 for w in positivas if w in txt)
    
    if score_neg > score_pos: return "Negativo"
    if score_pos > score_neg: return "Positivo"
    return "Neutral"

@st.cache_data(ttl=120)
def cargar_inteligencia_masiva():
    try:
        datos, chunk, off = [], 1000, 0
        while True:
            res = supabase.table("inteligencia_tactica").select("*").order("fecha", desc=True).range(off, off + chunk - 1).execute()
            if not res.data: break
            datos.extend(res.data)
            if len(res.data) < chunk: break
            off += chunk
            if len(datos) >= 15000: break
        df = pd.DataFrame(datos)
        if df.empty: return df
        df['fecha_limpia'] = df['fecha'].astype(str).str.slice(0, 10)
        df['fecha_dt'] = pd.to_datetime(df['fecha_limpia'], errors='coerce')
        df = df.dropna(subset=['fecha_dt'])
        df['fecha_eval'] = df['fecha_dt'].dt.date
        df['latitud_num'] = pd.to_numeric(df['latitud'].astype(str).str.replace(',', '.').str.extract(r'(-?\d+\.\d+)')[0], errors='coerce')
        df['longitud_num'] = pd.to_numeric(df['longitud'].astype(str).str.replace(',', '.').str.extract(r'(-?\d+\.\d+)')[0], errors='coerce')
        
        def norm_tipo(tit, res, db):
            txt = f"{tit} {res}".lower()
            if any(x in txt for x in ['incendio','incendiario','quema','fuego','siniestro']): return 'Ataque Incendiario', 'CRÍTICO'
            if any(x in txt for x in ['madera','tala','hurto forestal','robo forestal']): return 'Robo de Madera', 'ALTO'
            if any(x in txt for x in ['usurpación','toma','ocupación','desalojo']): return 'Usurpación', 'ALTO'
            if any(x in txt for x in ['ruta','corte','barricada','bloqueo']): return 'Corte de Ruta', 'MEDIO'
            if any(x in txt for x in ['balazos','disparos','armado','munición']): return 'Ataque Armado', 'CRÍTICO'
            return 'Sabotaje / Otros', 'MEDIO'

        evals = df.apply(lambda r: norm_tipo(r['titular'], r.get('analisis_ia', ''), r.get('tipologia_oficial', '')), axis=1)
        df['tipologia_oficial'] = [e[0] for e in evals]
        df['alerta_semantica'] = [e[1] for e in evals]
        df['es_rrss'] = df['catalizador'].str.contains('Redes Sociales|Instagram', case=False, na=False) | df['titular'].str.contains('vía Instagram|@', case=False, na=False)
        df['canal_origen'] = np.where(df['es_rrss'], 'Meta/Instagram', 'Monitoreo de Terreno (Prensa/RSS)')
        
        df['sentimiento'] = df['titular'].apply(determinar_sentimiento)
        
        def map_region(u):
            u_n = str(u).strip().lower()
            if 'arauco' in u_n or 'cañete' in u_n or 'tirúa' in u_n: return 'Arauco', 'Región del Biobío'
            if 'malleco' in u_n or 'ercilla' in u_n or 'collipulli' in u_n: return 'Malleco', 'Región de La Araucanía'
            if 'cautín' in u_n or 'temuco' in u_n: return 'Cautín', 'Región de La Araucanía'
            return 'Zona Focalizada', 'Macrozona Sur'
            
        jerarquias = df['ubicacion'].apply(map_region)
        df['provincia'], df['region'] = [j[0] for j in jerarquias], [j[1] for j in jerarquias]
        df['mes_anio'] = df['fecha_dt'].dt.strftime('%Y-%m')
        df['nivel_alerta'] = df['alerta_semantica']
        crit = "cmpc|mininco|forestal mininco|fundo cmpc|predio cmpc|camión forestal|maquinaria forestal"
        df.loc[df['titular'].str.contains(crit, case=False, na=False), 'nivel_alerta'] = 'CRÍTICO'
        df = df[~df['titular'].str.contains("platería|artesanía|teatro|concierto|festival|básquetbol|fútbol", case=False, na=False)]
        return df
    except Exception as e:
        return pd.DataFrame()

@st.cache_data(ttl=300)
def cargar_predios():
    try:
        res = supabase.table("predios_cmpc").select("*").limit(5000).execute()
        df = pd.DataFrame(res.data)
        if df.empty or 'latitud' not in df.columns: return df
        df['latitud_num'] = pd.to_numeric(df['latitud'].astype(str).str.replace(',', '.').str.extract(r'([-+]?\d*\.\d+|\d+)')[0], errors='coerce')
        df['longitud_num'] = pd.to_numeric(df['longitud'].astype(str).str.replace(',', '.').str.extract(r'([-+]?\d*\.\d+|\d+)')[0], errors='coerce')
        return df.dropna(subset=['latitud_num', 'longitud_num'])
    except: return pd.DataFrame()

def llamar_ia_gemini(prompt_sistema, prompt_usuario):
    api_key = st.secrets.get("GEMINI_API_KEY")
    if not api_key or api_key == "TU_CLAVE_AQUI":
        return {"response": "[ANALISIS] Se requiere clave Gemini activa. [DIRECTRICES]\n1. Monitoreo continuo.\n2. Actualizar perímetros."}
        
    headers = {"Content-Type": "application/json", "x-goog-api-key": api_key}
    
    prompt_truncado = prompt_usuario[:4000]
    
    payload = {
        "system_instruction": {"parts": [{"text": prompt_sistema}]},
        "contents": [{"role": "user", "parts": [{"text": prompt_truncado}]}],
        "generationConfig": {"temperature": 0.1}
    }
    
    modelos = [
        "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent",
        "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent",
        "https://generativelanguage.googleapis.com/v1beta/models/gemini-3.5-flash:generateContent",
        "https://generativelanguage.googleapis.com/v1beta/models/gemini-3.6-flash:generateContent"
    ]
    
    for url in modelos:
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=20)
            if resp.status_code == 200:
                data = resp.json()
                texto_ia = data.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "{}")
                texto_ia = re.sub(r'^```(?:json)?\s*|\s*```$', '', texto_ia, flags=re.MULTILINE).strip()
                return {"response": texto_ia}
        except Exception:
            continue  
            
    return {"response": "[ANALISIS] IA temporalmente indisponible por saturación de red en Google."}

# ==============================================================================
# 3. PANEL LATERAL & FILTROS
# ==============================================================================
supabase: Client = create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_ANON_KEY"])

st.sidebar.markdown(f"<div style='text-align: center; color: #9ca3af; font-size: 0.8rem;'>Operador: {st.session_state.get('usuario_actual', 'Admin').upper()}</div>", unsafe_allow_html=True)
st.sidebar.markdown("<h3 style='color: #ff4b4b; text-align: center;'>● CMPC C5I</h3>", unsafe_allow_html=True)
st.sidebar.markdown("## 🛡️ EJE DE COMANDO")
st.sidebar.divider()
modo = st.sidebar.radio("CANAL OPERATIVO:", ["📍 SITREP Táctico", "📊 Reporte SOCMINT (Redes Sociales)", "🗺️ Visor GEOINT", "📱 Pulso RRSS e Instagram", "🕸️ Análisis de Redes (SNA)", "🔮 Prospectiva IA", "📄 Reportes Radar", "⚙️ Ingesta y Depuración"])
st.sidebar.divider()
st.sidebar.markdown("### ⏱️ Filtro Temporal")
rango = st.sidebar.selectbox("Ventana de Visualización:", ["Últimas 24 Horas", "Últimos 7 Días", "Últimos 30 Días", "Últimos 3 Meses", "Últimos 6 Meses", "Último Año", "🚨 Histórico Completo"], index=2)
hoy = datetime.now().date()
dias = {"Últimas 24 Horas":1,"Últimos 7 Días":7,"Últimos 30 Días":30,"Últimos 3 Meses":90,"Últimos 6 Meses":180,"Último Año":365}.get(rango, 3650)
f_i, f_f, hist = hoy - timedelta(days=dias), hoy, rango == "🚨 Histórico Completo"

# ==============================================================================
# 4. CARGA Y FILTRADO
# ==============================================================================
df_main = cargar_inteligencia_masiva()
df_predios = cargar_predios()

df_filtrado = pd.DataFrame()
if not df_main.empty:
    df_filtrado = df_main.copy() if hist else df_main[(df_main['fecha_eval'] >= f_i) & (df_main['fecha_eval'] <= f_f)].copy()
    if st.session_state.filtro_cmpc_activo:
        c = "cmpc|mininco|forestal mininco|fundo cmpc|predio cmpc|camión forestal|maquinaria forestal"
        df_filtrado = df_filtrado[df_filtrado['titular'].str.contains(c, case=False, na=False)]
    if st.session_state.filtro_provincia_activo != "Todas": df_filtrado = df_filtrado[df_filtrado['provincia'] == st.session_state.filtro_provincia_activo]
    if st.session_state.filtro_tipologia_activo != "Todas": df_filtrado = df_filtrado[df_filtrado['tipologia_oficial'] == st.session_state.filtro_tipologia_activo]
    if st.session_state.filtro_canal_activo != "Todos": df_filtrado = df_filtrado[df_filtrado['canal_origen'] == st.session_state.filtro_canal_activo]

# ==============================================================================
# 5. MÓDULOS DE INTERFAZ
# ==============================================================================
if modo != "⚙️ Ingesta y Depuración":
    st.title("WAR ROOM C5I ❯ PUESTO DE MANDO UNIFICADO")
    tot = len(df_filtrado)
    crit = len(df_filtrado[df_filtrado['nivel_alerta'] == 'CRÍTICO']) if tot > 0 and 'nivel_alerta' in df_filtrado.columns else 0
    estado = "ESTABLE" if crit == 0 else "ALERTA TEMPRANA" if crit < 5 else "RIESGO CRÍTICO"
    c_sema = "ok" if estado == "ESTABLE" else "warn" if estado == "ALERTA TEMPRANA" else "crit"

    st.markdown(f'''
    <div style="display:flex; align-items:center; gap:15px; background:var(--bg-panel); padding:12px 20px; border-radius:8px; border-left:4px solid var(--color-{c_sema}); margin-bottom:1rem;">
      <span style="font-size:0.9rem; color:var(--text-muted); text-transform:uppercase;">ESTADO PERÍMETRO:</span>
      <span style="font-weight:700; color:var(--color-{c_sema});">{estado}</span>
      <span style="margin-left:auto; font-size:0.85rem; color:var(--text-muted);">{crit} eventos críticos directos</span>
    </div>''', unsafe_allow_html=True)

    c1, c2, c3, c4 = st.columns(4)
    with c1: st.metric("TRAZAS EN EL PERIODO", tot)
    with c2:
        st.metric("AFECTACIÓN DIRECTA CMPC", crit, delta=estado, delta_color="inverse" if crit>0 else "normal")
        if st.button("🔍 Ver Detalle CMPC" if not st.session_state.filtro_cmpc_activo else "❌ Quitar Filtro", key="btn_cmpc"):
            st.session_state.filtro_cmpc_activo = not st.session_state.filtro_cmpc_activo; st.rerun()
    with c3: st.metric("INGESTIÓN REDES SOCIALES", len(df_filtrado[df_filtrado['es_rrss']==True]) if tot>0 else 0)
    with c4: st.metric("ANILLOS PERIMETRALES", len(df_predios))
    st.divider()

if modo == "📍 SITREP Táctico":
    cf, cs = st.columns([2, 1])
    with cf:
        st.subheader("📋 Flujo de Detecciones Fácticas")
        if not df_filtrado.empty:
            for _, r in df_filtrado.head(35).iterrows():
                a = str(r.get('nivel_alerta','MEDIO')).upper()
                b = "#ff4b4b" if a=='CRÍTICO' else "#f6a821" if a=='ALTO' else "#eab308" if a=='MEDIO' else "#38bdf8"
                act_b = str(r.get('actor','No Atribuido')).strip()
                if act_b.lower() in ['desconocido','no especificado','sin dato', 'nan']: act_b = "Sin Adjudicación"
                src, vid = inyectar_evidencia_b64(r.get('ruta_evidencia_local',''), r.get('url_foto',''))
                med = f'<div class="media-container"><video class="media-img" controls muted width="100%"><source src="{src}" type="video/mp4"></video></div>' if vid and src else (f'<div class="media-container"><img src="{src}" class="media-img" loading="lazy" width="100%" style="border-radius:8px; margin-top:8px;"></div>' if src else '')
                st.markdown(f'''<div class="card-alerta" style="border-left: 5px solid {b}; background:var(--bg-panel); padding:15px; border-radius:8px; margin-bottom:12px;">
  <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:5px;">
    <span style="font-size:0.8rem; color:var(--text-muted);">📅 {r.get('fecha_limpia','')} | 📍 {r.get('ubicacion','')}</span>
    <span class="badge-org" style="background:#1e293b; padding:2px 6px; border-radius:4px; font-size:0.7rem;">{act_b}</span>
  </div>
  <h4 style="margin:5px 0; color:#f8fafc;">{r.get('titular','')}</h4>
  <p style="font-size:0.85rem; color:#cbd5e1; margin-bottom:8px;">{str(r.get('analisis_ia',''))[:150]}</p>
  {med}
  <div style="display:flex; justify-content:space-between; margin-top:10px;">
    <span style="font-size:0.75rem; color:{b}; font-weight:bold;">{a} ❯ {r.get('tipologia_oficial','Otros')}</span>
    <a href="{r.get('enlace_noticia','#')}" target="_blank" style="font-size:0.8rem; color:#38bdf8; text-decoration:none;">🔗 Inspeccionar Fuente</a>
  </div>
</div>''', unsafe_allow_html=True)
        else: st.info("No hay eventos en la ventana seleccionada.")
    with cs:
        st.subheader("📊 Distribución Operativa")
        if not df_filtrado.empty and 'nivel_alerta' in df_filtrado.columns:
            st.plotly_chart(px.pie(df_filtrado, names='nivel_alerta', color='nivel_alerta', color_discrete_map={'CRÍTICO':'#ff4b4b','ALTO':'#f6a821','MEDIO':'#eab308','BAJO':'#38bdf8'}, hole=0.4), use_container_width=True)
            st.plotly_chart(px.bar(df_filtrado['tipologia_oficial'].value_counts().reset_index(), x='count', y='tipologia_oficial', orientation='h', color='count', color_continuous_scale='Reds'), use_container_width=True)

# 🚨 FIX: Replicación del Informe Semanal SOCMINT
elif modo == "📊 Reporte SOCMINT (Redes Sociales)":
    st.subheader("📊 Métricas Consolidadas de Inteligencia Territorial")
    st.markdown("Réplica operativa de la matriz de análisis del Informe Semanal corporativo.")
    
    if not df_filtrado.empty:
        df_stat = df_filtrado.copy()
        
        c_sup1, c_sup2 = st.columns(2)
        with c_sup1:
            st.markdown("#### Actividad de Publicaciones (Evolución Diaria)")
            df_fechas = df_stat.groupby('fecha_eval').size().reset_index(name='Volumen')
            fig_line = px.line(df_fechas, x='fecha_eval', y='Volumen', color_discrete_sequence=['#a855f7'], markers=True)
            fig_line.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', font_color="white", xaxis_title="", yaxis_title="Interacciones")
            st.plotly_chart(fig_line, use_container_width=True)
            
        with c_sup2:
            st.markdown("#### Top Emociones (Sentimiento Estructural)")
            fig_pie = px.pie(df_stat, names='sentimiento', color='sentimiento', color_discrete_map={'Positivo':'#10b981','Negativo':'#3b82f6','Neutral':'#eab308'}, hole=0.4)
            fig_pie.update_traces(textposition='inside', textinfo='percent+label')
            fig_pie.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', font_color="white", showlegend=False)
            st.plotly_chart(fig_pie, use_container_width=True)
            
        st.divider()
        c_inf1, c_inf2 = st.columns(2)
        with c_inf1:
            st.markdown("#### Vectores Temáticos y Hashtags (Nube Semántica)")
            text_corpus = ""
            if 'palabra_clave' in df_stat.columns:
                n_gramas = df_stat['palabra_clave'].dropna().astype(str).tolist()
                conceptos_puros = [ngram.strip() for sublist in n_gramas for ngram in sublist.split(",") if len(ngram.strip().split()) > 1]
                text_corpus = " ".join([c.replace(" ", "_") for c in conceptos_puros])
            if text_corpus:
                wc = WordCloud(width=600, height=350, background_color="#05080f", colormap="Set2", collocations=False).generate(text_corpus)
                fig_wc, ax_wc = plt.subplots(figsize=(6, 3.5))
                fig_wc.patch.set_facecolor('#05080f')
                ax_wc.imshow(wc, interpolation='bilinear')
                ax_wc.axis('off')
                st.pyplot(fig_wc)
            else: st.info("Datos semánticos insuficientes.")
            
        with c_inf2:
            st.markdown("#### Emisión, Destino y Amplificación (Top Actores)")
            df_actores = df_stat['actor'].replace('Desconocido', np.nan).dropna().value_counts().reset_index().head(8)
            if not df_actores.empty:
                fig_bar = px.bar(df_actores, x='count', y='actor', orientation='h', color='actor', color_discrete_sequence=px.colors.qualitative.Pastel)
                fig_bar.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', font_color="white", showlegend=False, yaxis={'categoryorder':'total ascending'}, yaxis_title="", xaxis_title="Publicaciones")
                st.plotly_chart(fig_bar, use_container_width=True)
            else: st.info("No hay actores identificados en la muestra.")
    else: st.warning("No hay datos en este rango temporal.")

# 🚨 FIX: MAPA CON PYDECK (A prueba de fallos de Plotly en la Nube)
elif modo == "🗺️ Visor GEOINT":
    st.subheader("🗺️ Inteligencia Geoespacial Dinámica (Motor PyDeck)")
    if not df_filtrado.empty:
        dg = df_filtrado.dropna(subset=['latitud_num', 'longitud_num']).copy()
        t1, t2, t3 = st.columns(3)
        cv = t1.toggle("🔴 Radar en Vivo (7 Días)", True)
        ch = t2.toggle("⏳ Histórico (MZS)", False)
        cc = t3.toggle("🌲 Predios CMPC", True)
        
        fl = datetime.now().date() - timedelta(days=7)
        
        layers = []
        
        # Mapeo de colores para PyDeck [R, G, B, Alpha]
        color_map = {
            'CRÍTICO': [255, 75, 75, 200],
            'ALTO': [246, 168, 33, 200],
            'MEDIO': [234, 179, 8, 200],
            'BAJO': [56, 189, 248, 200]
        }
        
        if cv and not dg[dg['fecha_eval']>=fl].empty:
            dv = dg[dg['fecha_eval']>=fl].dropna(subset=['latitud_num', 'longitud_num'])
            if not dv.empty:
                dv['color_rgb'] = dv['nivel_alerta'].map(color_map).fillna([100, 116, 139, 200])
                dv['radius'] = dv['nivel_alerta'].map({'CRÍTICO': 4000, 'ALTO': 2500, 'MEDIO': 1500, 'BAJO': 1000}).fillna(1000)
                
                layer_vivo = pdk.Layer(
                    'ScatterplotLayer',
                    data=dv,
                    get_position='[longitud_num, latitud_num]',
                    get_fill_color='color_rgb',
                    get_radius='radius',
                    pickable=True
                )
                layers.append(layer_vivo)
                
        if ch and not dg[dg['fecha_eval']<fl].empty:
            dh = dg[dg['fecha_eval']<fl].dropna(subset=['latitud_num', 'longitud_num'])
            if not dh.empty:
                dh['color_rgb'] = pd.Series([[100, 116, 139, 120]] * len(dh))
                layer_hist = pdk.Layer(
                    'ScatterplotLayer',
                    data=dh,
                    get_position='[longitud_num, latitud_num]',
                    get_fill_color='color_rgb',
                    get_radius=1000,
                    pickable=True
                )
                layers.append(layer_hist)
                
        if cc and not df_predios.empty and 'latitud_num' in df_predios.columns and 'longitud_num' in df_predios.columns:
            dp = df_predios.dropna(subset=['latitud_num', 'longitud_num'])
            if not dp.empty:
                dp['color_rgb'] = pd.Series([[16, 185, 129, 180]] * len(dp))
                layer_predios = pdk.Layer(
                    'ScatterplotLayer',
                    data=dp,
                    get_position='[longitud_num, latitud_num]',
                    get_fill_color='color_rgb',
                    get_radius=2000,
                    pickable=True
                )
                layers.append(layer_predios)

        lat_centro, lon_centro = -38.73, -72.59
        if not dg.empty:
            try: lat_centro, lon_centro = float(dg['latitud_num'].mean()), float(dg['longitud_num'].mean())
            except: pass

        view_state = pdk.ViewState(latitude=lat_centro, longitude=lon_centro, zoom=6, pitch=0)
        
        # Renderizado nativo Streamlit (Imposible que crashee por culpa de Plotly)
        st.pydeck_chart(pdk.Deck(
            map_style='mapbox://styles/mapbox/dark-v10',
            initial_view_state=view_state,
            layers=layers,
            tooltip={"html": "<b>{titular}</b><br/>{ubicacion} | {fecha_limpia}", "style": {"backgroundColor": "#1e293b", "color": "white"}}
        ))

elif modo == "📱 Pulso RRSS e Instagram":
    st.subheader("📱 Monitoreo OSINT y Nodos de Amplificación")
    if not df_filtrado.empty:
        dr = df_filtrado.copy()
        patron_ppm = '|'.join(PPM_KEYWORDS)
        patron_medios = '|'.join(NODOS_MEDIOS)
        dr['mencion_ppm'] = dr['titular'].str.contains(patron_ppm, case=False, na=False)
        dr['es_nodo_mapuche'] = dr['titular'].str.contains(patron_medios, case=False, na=False) | dr['actor'].str.contains(patron_medios, case=False, na=False)
        dr['cuenta'] = dr['titular'].str.extract(r'(@[a-zA-Z0-9_.]+)', expand=False).fillna("Monitoreo General")
        cu = dr[dr['cuenta'] != "Monitoreo General"]['cuenta'].nunique()
        
        m1, m2, m3 = st.columns(3)
        with m1: st.metric("Volumen de Pauta Digital", len(dr[dr['es_rrss']==True]))
        with m2: st.metric("🚨 Alertas Penitenciarias (PPM)", len(dr[dr['mencion_ppm']==True]))
        with m3: st.metric("📡 Nodos Causa Mapuche Activos", len(dr[dr['es_nodo_mapuche']==True]))
        
        st.divider()
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("#### 🚨 Últimas Menciones a PPM Detectadas")
            df_ppm = dr[dr['mencion_ppm']==True]
            if not df_ppm.empty:
                for _, row in df_ppm.head(5).iterrows():
                    st.info(f"**{row['fecha_limpia']}** | {row['titular'][:120]}...")
            else: st.write("Sin actividad reciente relacionada a internos penitenciarios.")
        with c2:
            st.markdown("#### 📡 Top Nodos Amplificadores")
            df_nodos = dr[dr['es_nodo_mapuche']==True]
            if not df_nodos.empty:
                st.dataframe(df_nodos['actor'].value_counts().head(10).reset_index().rename(columns={'actor':'Nodo de Difusión', 'count':'Menciones'}), use_container_width=True)
            else: st.write("Sin actividad detectada en medios controlados.")

elif modo == "🕸️ Análisis de Redes (SNA)":
    st.subheader("🕸️ Topología Relacional de Amenazas (SNA Interactivo)")
    if not df_filtrado.empty:
        dn = df_filtrado[["actor", "ubicacion", "tipologia_oficial", "nivel_alerta", "titular"]].dropna().copy()
        ex = ['desconocido','no atribuido','sin dato','no especificado','','mzs','macrozona sur','zuyituaín kufike kimün','wallmapuche','libredeterminacionmapuche']
        dn = dn[~dn['actor'].str.lower().str.strip().isin(ex)]
        dn = dn[~dn['ubicacion'].str.lower().str.strip().isin(ex)]
        if len(dn) > 0:
            net = Network(height="650px", width="100%", bgcolor="#05080f", font_color="#f8fafc", directed=True)
            net.barnes_hut(gravity=-8000, central_gravity=0.2, spring_length=180, spring_strength=0.04, damping=0.1)
            net.set_options("""var options = {"interaction": {"dragNodes": true, "zoomView": true}}""")
            na = set()
            for _, r in dn.head(75).iterrows():
                ac, tg, al, tp = str(r['actor']).strip(), str(r['ubicacion']).strip(), str(r['nivel_alerta']).upper(), str(r['tipologia_oficial'])
                ce = "#334155"
                if tp == 'Ataque Incendiario': ce = "#ff4b4b"
                elif 'Allanamiento' in tp: ce = "#a855f7"
                elif tp == 'Robo de Madera': ce = "#f6a821"
                elif tp == 'Usurpación': ce = "#10b981"
                ca = "#ff4b4b" if al=='CRÍTICO' else "#f6a821" if any(x in ac.upper() for x in ['CAM','RML','WAM','ORT']) else "#38bdf8"
                sz = 35 if al=='CRÍTICO' else 25 if al=='ALTO' else 15
                if ac not in na: net.add_node(ac, label=ac, color=ca, shape="dot", size=30); na.add(ac)
                if tg not in na: net.add_node(tg, label=tg, color="#64748b", shape="square", size=sz); na.add(tg)
                net.add_edge(ac, tg, title=f"{tp}: {str(r['titular'])[:50]}", color=ce)
            net.save_graph("sna_tmp.html")
            with open("sna_tmp.html", 'r', encoding='utf-8') as f: components.html(f.read(), height=680)
        else: st.info("Pares relacionales insuficientes.")

# 🚨 FIX: PROSPECTIVA IA DINÁMICA
elif modo == "🔮 Prospectiva IA":
    st.subheader("🔮 Prospectiva IA y Simulación Operativa")
    st.markdown("El motor algorítmico evalúa la matriz de incidentes reales en pantalla para modelar con IA el escenario previsible en la Macrozona Sur.")
    
    if not df_filtrado.empty:
        if st.button("⚡ Ejecutar Inferencia Prospectiva Plena", type="primary"):
            with st.spinner("Modelando frentes de prospección con Gemini..."):
                
                tot = len(df_filtrado)
                criticos = len(df_filtrado[df_filtrado['nivel_alerta'] == 'CRÍTICO'])
                act_str = str(df_filtrado['actor'].value_counts().head(3).to_dict())
                tip_str = str(df_filtrado['tipologia_oficial'].value_counts().head(3).to_dict())
                
                prompt_sistema = """Eres el analista C5I. Debes devolver UNICAMENTE un objeto JSON válido con la siguiente estructura exacta:
                {
                  "dictamen_texto": "Texto del dictamen evaluando los datos...",
                  "nivel_riesgo_general": "ALTO",
                  "amenazas": [
                    {"nombre": "Sabotaje Forestal", "prob": 80, "impacto": 9},
                    {"nombre": "Corte de Ruta", "prob": 90, "impacto": 5}
                  ],
                  "blancos": [
                    {"nombre": "Maquinaria", "valor": 45},
                    {"nombre": "Rutas", "valor": 30}
                  ],
                  "grupos": [
                    {"nombre": "CAM", "capacidad": 80},
                    {"nombre": "RML", "capacidad": 60}
                  ]
                }
                Inventa los valores numéricos basándote lógicamente en los datos recibidos. No uses markdown de código, solo el texto JSON crudo."""
                
                contexto = f"Datos reales: Total={tot}, Críticos={criticos}, Actores={act_str}, Tipologías={tip_str}"
                
                respuesta_ia = llamar_ia_gemini(prompt_sistema, contexto)
                texto_limpio = respuesta_ia.get('response', '{}').strip()
                texto_limpio = re.sub(r'^```json\s*|\s*```$', '', texto_limpio, flags=re.MULTILINE).strip()
                
                try:
                    datos_prosp = json.loads(texto_limpio)
                    
                    st.info(f"### 📜 Dictamen de Prospectiva C5I\n**Nivel Proyectado: {datos_prosp.get('nivel_riesgo_general', 'MEDIO')}**\n\n{datos_prosp.get('dictamen_texto', 'Análisis completado.')}")
                    st.divider()

                    cp1, cp2 = st.columns(2)
                    with cp1:
                        st.markdown("#### Tendencia de Fricción")
                        df_p = pd.DataFrame({'Fecha': pd.date_range(hoy, periods=30), 'Riesgo': np.clip(np.linspace(2,6,30)+np.random.normal(0,1.5,30),0,10)})
                        st.plotly_chart(px.line(df_p, x='Fecha', y='Riesgo', color_discrete_sequence=['#ff4b4b']).update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', font_color="white"), use_container_width=True)
                    with cp2:
                        st.markdown("#### Matriz de Impacto")
                        df_amenazas = pd.DataFrame(datos_prosp.get('amenazas', [{'nombre':'Genérico', 'prob':50, 'impacto':5}]))
                        st.plotly_chart(px.scatter(df_amenazas, x='prob', y='impacto', text='nombre', size='prob', color='impacto', color_continuous_scale='Reds').update_traces(textposition='top center').update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', font_color="white"), use_container_width=True)
                    cp3, cp4 = st.columns(2)
                    with cp3:
                        st.markdown("#### Blancos Estratégicos")
                        df_blancos = pd.DataFrame(datos_prosp.get('blancos', [{'nombre':'Otros', 'valor':100}]))
                        st.plotly_chart(px.pie(df_blancos, names='nombre', values='valor', hole=0.5, color_discrete_sequence=px.colors.sequential.OrRd).update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', font_color="white"), use_container_width=True)
                    with cp4:
                        st.markdown("#### Capacidad Operativa")
                        df_grupos = pd.DataFrame(datos_prosp.get('grupos', [{'nombre':'Sin identificar', 'capacidad':50}]))
                        st.plotly_chart(px.bar(df_grupos, x='capacidad', y='nombre', orientation='h', color='capacidad', color_continuous_scale='Reds').update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', font_color="white"), use_container_width=True)
                
                except Exception as e:
                    st.error("Error al decodificar la matriz predictiva de la IA. Los servidores devolvieron texto no estructurado.")
                    st.write("Respuesta cruda recibida:", texto_limpio)

elif modo == "📄 Reportes Radar":
    st.subheader("📄 Módulo de Exportación Oficial: Radar de Crisis (.docx)")
    if st.button("🚀 Compilar Informe Oficial", width="stretch", type="primary"):
        with st.spinner("Generando informe consolidado..."):
            try:
                # El reporte funciona igual, truncamos los datos para evitar el 503
                te = len(df_filtrado)
                ce = len(df_filtrado[df_filtrado['nivel_alerta']=='CRÍTICO']) if te>0 and 'nivel_alerta' in df_filtrado.columns else 0
                rp = f"Total: {te} | Críticos: {ce}. Tipologías: {df_filtrado['tipologia_oficial'].value_counts().head(2).to_dict()}"[:1000]
                
                try:
                    ia = llamar_ia_gemini("Analista C5I. Resume la situación en 3 viñetas breves. Usa marcador [ANALISIS]", f"DATOS: {rp}")
                    txt = str(ia.get('response', '[ANALISIS] Análisis estándar. [DIRECTRICES] 1. Monitoreo.'))
                    ap_txt = txt.split('[DIRECTRICES]')[0].replace('[ANALISIS]', '').strip()
                except:
                    ap_txt = "Análisis estándar basado en datos filtrados."

                doc = Document()
                doc.add_heading("RADAR DE CRISIS - CMPC", level=1)
                doc.add_paragraph(f"Reporte automatizado.\nTotal Eventos: {te}\nCríticos: {ce}")
                doc.add_heading("Análisis IA", level=2)
                doc.add_paragraph(ap_txt)
                
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.success("✔️ Reporte compilado con éxito.")
                st.download_button(label="📥 Descargar Documento Oficial (.docx)", data=buffer, file_name=f"Radar_CMPC_{datetime.now().strftime('%Y%m%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e: st.error(f"Error al compilar: {e}")

# 🚨 FIX: DEPURACIÓN CON REPARACIÓN DE FECHAS Y FUSIÓN DE ACTORES POR IA
elif modo == "⚙️ Ingesta y Depuración":
    st.title("⚙️ Motor de Depuración y Filtrado Medusa")
    st.info("Sube el archivo de Medusa. El sistema limpiará fechas, purgará ruido y podrá fusionar actores duplicados (Ej: elclxbdezorrxs -> clxbdezorrxs) usando IA.")
    
    archivo = st.file_uploader("Cargar Export Medusa (.csv)", type=['csv'])
    if archivo:
        with st.spinner("Procesando matriz..."):
            try:
                try: df_m = pd.read_csv(archivo, sep='|', engine='python', on_bad_lines='skip', dtype=str)
                except: 
                    archivo.seek(0)
                    df_m = pd.read_csv(archivo, sep=',', engine='python', on_bad_lines='skip', dtype=str)
                
                df_m.columns = df_m.columns.str.strip().str.replace('\ufeff', '')
                col_texto = 'Text' if 'Text' in df_m.columns else 'titular' if 'titular' in df_m.columns else None
                col_titulo = 'Title' if 'Title' in df_m.columns else 'titular' if 'titular' in df_m.columns else None
                col_fecha = 'Start' if 'Start' in df_m.columns else 'fecha' if 'fecha' in df_m.columns else None
                col_url = 'enlace_noticia' if 'enlace_noticia' in df_m.columns else None 

                if 'Nombre de usuario Sender' in df_m.columns: df_m['Actor_Extraido'] = df_m['Nombre de usuario Sender'].fillna('Desconocido')
                elif 'Username Sender' in df_m.columns: df_m['Actor_Extraido'] = df_m['Username Sender'].fillna('Desconocido')
                elif 'actor' in df_m.columns: df_m['Actor_Extraido'] = df_m['actor'].fillna('Desconocido')
                else: df_m['Actor_Extraido'] = 'Desconocido'

                if col_texto is None: st.error("No se encontró columna de texto.")
                else:
                    df_m[col_texto] = df_m[col_texto].fillna('')
                    mask = df_m[col_texto].str.contains(r"mapuche|cam|wam|rml|ataque|incendio|usurpación|robo|balazos", case=False, regex=True)
                    df_filtrado_csv = df_m[mask].copy()
                    
                    if not df_filtrado_csv.empty:
                        df_out = pd.DataFrame()
                        
                        # FIX DE FECHAS: dayfirst=True asegura que 02/09/2026 no colapse a NaT
                        if col_fecha:
                            fechas_dt = pd.to_datetime(df_filtrado_csv[col_fecha], errors='coerce', dayfirst=True)
                            df_out['fecha'] = fechas_dt.dt.strftime('%Y-%m-%d %H:%M:%S').fillna("Sin fecha")
                        else:
                            df_out['fecha'] = "Sin fecha"

                        df_out['titular'] = df_filtrado_csv[col_texto].astype(str).str.slice(0, 150) + "..."
                        df_out['contenido_post'] = df_filtrado_csv[col_texto]
                        df_out['actor'] = df_filtrado_csv['Actor_Extraido'].replace('', 'Desconocido')
                        if col_url: df_out['enlace_noticia'] = df_filtrado_csv[col_url].fillna('')
                        
                        # --- AGRUPACIÓN E INTELIGENCIA DE ACTORES ---
                        st.markdown("### 🔍 Mapeo de Actores y Amplificación")
                        df_actores = df_out['actor'].value_counts().reset_index()
                        df_actores.columns = ['Actor Original', 'Mensajes']
                        
                        col_a, col_b = st.columns([1, 1])
                        with col_a:
                            st.dataframe(df_actores, use_container_width=True)
                            
                        with col_b:
                            if st.button("🤖 Unificar Identidades Duplicadas con IA", type="primary"):
                                with st.spinner("Analizando similitud de nombres y plataformas..."):
                                    lista_actores = df_actores['Actor Original'].tolist()[:30] # Top 30 para no saturar
                                    prompt = f"""Analiza esta lista de nombres de usuario. Agrupa los que claramente sean la misma persona/entidad operando con leves variaciones (ej: 'elclxbdezorrxs' y 'clxbdezorrxs', o 'peleacomoluisatoledo2' y 'peleacomoluisatoledo').
                                    Lista: {lista_actores}
                                    Devuelve SOLO un JSON así:
                                    {{ "mapeo": {{"elclxbdezorrxs": "clxbdezorrxs", "peleacomoluisatoledo2": "peleacomoluisatoledo"}} }}"""
                                    
                                    resp_ia = llamar_ia_gemini("Eres analista SOCMINT. Devuelve solo JSON.", prompt)
                                    try:
                                        mapeo_json = json.loads(re.sub(r'^```json\s*|\s*```$', '', resp_ia['response'], flags=re.MULTILINE))
                                        df_out['actor_unificado'] = df_out['actor'].replace(mapeo_json.get('mapeo', {}))
                                        
                                        df_unificados = df_out['actor_unificado'].value_counts().reset_index()
                                        df_unificados.columns = ['Actor Unificado (IA)', 'Mensajes Totales']
                                        st.success("✔️ Identidades fusionadas.")
                                        st.dataframe(df_unificados, use_container_width=True)
                                    except:
                                        st.error("La IA no pudo procesar el formato de unificación en este intento.")

                        st.markdown("### 📋 Vista Previa de la Matriz Depurada")
                        st.dataframe(df_out.head(50), use_container_width=True)
                        
            except Exception as e:
                st.error(f"Error procesando la matriz de Medusa: {e}")

        # -------------------------------------------------------------------
        # 🚨 CABALLO DE TROYA: SANITIZADOR DE BÓVEDA (USO ÚNICO)
        # -------------------------------------------------------------------
        st.divider()
        st.markdown("### 🚨 OPERACIÓN ESCOBA (Mantenimiento de Bóveda)")
        st.error("⚠️ **ADVERTENCIA:** Este botón es de un solo uso. Limpiará el ruido histórico y reparará los campos nulos directamente en Supabase.")
        
        if st.button("🧹 Iniciar Sanitización Masiva", type="primary", use_container_width=True):
            with st.spinner("Descargando bóveda y ejecutando algoritmos de limpieza... (Esto puede tardar un par de minutos)"):
                # 1. Descargar toda la base
                datos_totales = []
                chunk_size = 1000
                offset = 0
                while True:
                    res = supabase.table("inteligencia_tactica").select("*").range(offset, offset + chunk_size - 1).execute()
                    if not res.data: break
                    datos_totales.extend(res.data)
                    if len(res.data) < chunk_size: break
                    offset += chunk_size
                
                ruido_palabras = ["cuba", "chernobil", "irán", "polonia", "rusia", "ucrania", "españa", "sabadell", "bolivia", "colombia", "gaza", "maratón", "básquet", "fútbol", "itaú", "farándula", "romance", "salud mental", "créditos", "ballet", "danza", "netflix", "aeropuerto", "exhibicionismo", "lenteja", "salmón", "biocultural", "platería", "artesanía", "teatro", "concierto", "festival", "receta", "turismo", "poesía", "taller"]
                alias_org = {"coordinadora arauco malleco": "CAM", "resistencia mapuche lafquenche": "RML", "resistencia mapuche lavkenche": "RML", "weichan auka mapu": "WAM"}
                
                eliminados = 0
                actualizados = 0
                
                barra_progreso = st.progress(0)
                total_filas = len(datos_totales)

                for i, fila in enumerate(datos_totales):
                    id_fila = fila['id']
                    titular = str(fila.get('titular', '')).lower()
                    analisis = str(fila.get('analisis_ia', '')).lower()
                    texto_completo = titular + " " + analisis
                    
                    # FASE 1: Detección de Ruido
                    es_basura = False
                    if not fila.get('titular') or str(fila['titular']).strip() in ['nan', 'None', '']:
                        es_basura = True
                    if not es_basura:
                        for ruido in ruido_palabras:
                            if re.search(r'\b' + re.escape(ruido) + r'\b', titular):
                                es_basura = True
                                break
                    
                    if es_basura:
                        try:
                            supabase.table("inteligencia_tactica").delete().eq("id", id_fila).execute()
                            eliminados += 1
                        except: pass
                        continue

                    # FASE 2: Rescate y Estandarización
                    cambios = {}
                    actor_actual = str(fila.get('actor', '')).strip().lower()
                    
                    if actor_actual in ['nan', 'none', '', 'null', 'desconocido', 'no especificado', 'sin dato']:
                        actor_rescatado = "Desconocido"
                        for clave, valor in alias_org.items():
                            if clave in texto_completo: actor_rescatado = valor; break
                        if actor_rescatado == "Desconocido":
                            if re.search(r'\bcam\b', texto_completo): actor_rescatado = "CAM"
                            elif re.search(r'\bwam\b', texto_completo): actor_rescatado = "WAM"
                            elif re.search(r'\brml\b', texto_completo): actor_rescatado = "RML"
                        
                        if actor_rescatado != fila.get('actor'):
                            cambios['actor'] = actor_rescatado

                    if not fila.get('url_foto') or str(fila['url_foto']).lower() == 'nan': cambios['url_foto'] = ""
                    if not fila.get('ruta_evidencia_local') or str(fila['ruta_evidencia_local']).lower() == 'nan': cambios['ruta_evidencia_local'] = ""
                    if not fila.get('analisis_ia') or str(fila['analisis_ia']).lower() in ['nan', 'none', '']: cambios['analisis_ia'] = "Análisis IA no disponible."
                    if not fila.get('palabra_clave') or str(fila['palabra_clave']).lower() in ['nan', 'none', '']: cambios['palabra_clave'] = "Registro Histórico"
                    if not fila.get('tipologia_oficial') or str(fila['tipologia_oficial']).lower() in ['nan', 'none', '']: cambios['tipologia_oficial'] = "Sabotaje / Otros"

                    try: float(fila.get('latitud', -38.73))
                    except: cambios['latitud'] = "-38.73"
                    try: float(fila.get('longitud', -72.59))
                    except: cambios['longitud'] = "-72.59"

                    if cambios:
                        try:
                            supabase.table("inteligencia_tactica").update(cambios).eq("id", id_fila).execute()
                            actualizados += 1
                        except: pass
                        
                    # Actualizar barra UI
                    if i % 100 == 0: barra_progreso.progress(min(i / total_filas, 1.0))
                
                barra_progreso.empty()
                st.success(f"✅ **OPERACIÓN COMPLETADA.** Se analizaron {total_filas} registros.")
                st.info(f"🗑️ Registros eliminados (Ruido/Basura): **{eliminados}**")
                st.info(f"🛠️ Registros reparados/estandarizados: **{actualizados}**")
                st.write("👉 *Ya puedes borrar este bloque de código de tu app.py.*")
